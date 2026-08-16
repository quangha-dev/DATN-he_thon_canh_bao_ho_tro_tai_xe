import json
import re
import sys
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "087F73"
MUTED = "5D6673"
LIGHT = "F4F6F9"
BORDER = "C8D1DC"
WHITE = "FFFFFF"
PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)


def compute_widths(rows):
    columns = len(rows[0])
    max_lengths = []
    for col in range(columns):
        length = max(len(re.sub(r"[`*_]", "", row[col])) for row in rows)
        max_lengths.append(max(5, min(length, 42)))

    if columns == 2 and max_lengths[0] <= 22:
        return [1900, PAGE_WIDTH_DXA - 1900]
    if columns == 3:
        return [1550, 2600, PAGE_WIDTH_DXA - 4150]
    if columns == 4:
        return [1050, 2200, 1500, PAGE_WIDTH_DXA - 4750]

    minimum = 850
    remaining = PAGE_WIDTH_DXA - minimum * columns
    total = sum(max_lengths)
    widths = [minimum + round(remaining * value / total) for value in max_lengths]
    widths[-1] += PAGE_WIDTH_DXA - sum(widths)
    return widths


def clean_markdown(text):
    text = text.strip()
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(.*?)\*(?!\*)", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text.replace("  ", " ")


INLINE_PATTERN = re.compile(
    r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|(?<!\*)\*[^*]+\*(?!\*))"
)


def add_inline(paragraph, text, size=11, base_color=None):
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=size, color=base_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=base_color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=max(8.5, size - 0.5), color=DARK_BLUE)
        elif token.startswith("["):
            link_match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            label, url = link_match.groups()
            run = paragraph.add_run(f"{label} ({url})")
            set_run_font(run, size=size, color=BLUE)
            run.underline = True
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, color=base_color, italic=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=base_color)


def paragraph_border_bottom(paragraph, color=BORDER, size="6", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Nhấn Ctrl+A, F9 trong Word nếu mục lục chưa được cập nhật."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, placeholder, fld_end])


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
        "Heading 4": (11, NAVY, 7, 3),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    code_style = styles.add_style("SafeFleet Code", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(8.5)
    code_style.paragraph_format.left_indent = Inches(0.2)
    code_style.paragraph_format.right_indent = Inches(0.2)
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.line_spacing = 1.0

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(10)


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run("SAFEFLEET  |  BÁO CÁO ĐỒ ÁN TỐT NGHIỆP")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    paragraph_border_bottom(hp, color="D8DEE7", size="4", space="4")

    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_cover(doc, title, subtitle, metadata):
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(16)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("BÁO CÁO ĐỒ ÁN TỐT NGHIỆP")
    set_run_font(run, size=12, color=TEAL, bold=True)
    kicker.paragraph_format.space_after = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, size=28, color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    set_run_font(run, size=14, color=DARK_BLUE, bold=True)
    p.paragraph_format.space_after = Pt(34)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Hệ thống SafeFleet")
    set_run_font(run, size=13, color=TEAL, bold=True)
    p.paragraph_format.space_after = Pt(24)

    for line in metadata:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(p, clean_markdown(line), size=10, base_color=MUTED)
        p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    run = p.add_run("Tháng 8 năm 2026")
    set_run_font(run, size=11, color=NAVY, bold=True)
    p.add_run().add_break(WD_BREAK.PAGE)

    toc_heading = doc.add_paragraph("MỤC LỤC", style="Heading 1")
    toc_heading.paragraph_format.page_break_before = False
    toc = doc.add_paragraph()
    add_toc(toc)
    toc.add_run().add_break(WD_BREAK.PAGE)


def add_table(doc, rows):
    if not rows:
        return
    columns = len(rows[0])
    rows = [row + [""] * (columns - len(row)) for row in rows]
    widths = compute_widths(rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_index == 0 or len(value) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, value, size=8.5, base_color=NAVY if row_index == 0 else None)
            for run in p.runs:
                if row_index == 0:
                    run.bold = True
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(3)


def add_image(doc, alt, png_path):
    with Image.open(png_path) as image:
        width_px, height_px = image.size
    ratio = height_px / max(width_px, 1)
    max_width = 6.35
    # Named figure-layout override: reserve room for the figure heading,
    # explanatory paragraph and caption on the same page.
    max_height = 6.95
    width = max_width
    height = width * ratio
    if height > max_height:
        height = max_height
        width = height / ratio

    if doc.paragraphs:
        doc.paragraphs[-1].paragraph_format.keep_with_next = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(png_path), width=Inches(width), height=Inches(height))
    try:
        shape._inline.docPr.set("descr", clean_markdown(alt))
    except Exception:
        pass

    caption = doc.add_paragraph(style="Caption")
    caption.add_run(clean_markdown(alt))


def parse_table_line(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_separator(line):
    cells = parse_table_line(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def build_document(markdown_path, manifest_path, output_path):
    markdown_path = Path(markdown_path).resolve()
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    manifest = json.loads(Path(manifest_path).read_text(encoding="utf-8"))

    title = clean_markdown(lines[0].lstrip("# "))
    subtitle = clean_markdown(lines[2].lstrip("# "))
    metadata = []
    for line in lines[3:]:
        if line.startswith(">"):
            metadata.append(line.lstrip("> ").rstrip())
        elif metadata and line.strip() == "---":
            break

    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    add_cover(doc, title, subtitle, metadata)

    start = next(index for index, line in enumerate(lines) if line.strip() == "## Tóm tắt")
    index = start
    paragraph_buffer = []
    in_code = False
    in_math = False
    code_lines = []
    math_lines = []

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            text = " ".join(part.strip() for part in paragraph_buffer if part.strip())
            if text:
                p = doc.add_paragraph()
                add_inline(p, text)
            paragraph_buffer = []

    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()

        if in_code:
            if line.startswith("```"):
                for code_line in code_lines or [""]:
                    p = doc.add_paragraph(style="SafeFleet Code")
                    p.paragraph_format.keep_together = True
                    run = p.add_run(code_line)
                    set_run_font(run, name="Consolas", size=8.5, color=NAVY)
                    p_pr = p._p.get_or_add_pPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), "F3F5F7")
                    p_pr.append(shd)
                doc.add_paragraph().paragraph_format.space_after = Pt(3)
                code_lines = []
                in_code = False
            else:
                code_lines.append(line)
            index += 1
            continue

        if in_math:
            if line.strip() == r"\]":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(" ".join(math_lines))
                set_run_font(run, name="Cambria Math", size=11, color=NAVY, italic=True)
                math_lines = []
                in_math = False
            else:
                math_lines.append(line.strip())
            index += 1
            continue

        if line.startswith("```"):
            flush_paragraph()
            in_code = True
            index += 1
            continue

        if line.strip() == r"\[":
            flush_paragraph()
            in_math = True
            index += 1
            continue

        image_match = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+\.svg)\)", line.strip())
        if image_match:
            flush_paragraph()
            alt, ref = image_match.groups()
            png_path = Path(manifest[ref])
            add_image(doc, alt, png_path)
            index += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            flush_paragraph()
            level = min(len(heading_match.group(1)), 4)
            heading_text = clean_markdown(heading_match.group(2))
            p = doc.add_paragraph(heading_text, style=f"Heading {level}")
            if level == 1:
                p.paragraph_format.page_break_before = True
            index += 1
            continue

        if line.strip().startswith("|") and line.strip().endswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            flush_paragraph()
            table_rows = [parse_table_line(line)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|") and lines[index].strip().endswith("|"):
                table_rows.append(parse_table_line(lines[index]))
                index += 1
            add_table(doc, table_rows)
            continue

        bullet_match = re.match(r"^\s*-\s+(.+)$", line)
        if bullet_match:
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, bullet_match.group(1))
            index += 1
            continue

        numbered_match = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if numbered_match:
            flush_paragraph()
            p = doc.add_paragraph(style="List Number")
            add_inline(p, numbered_match.group(1))
            index += 1
            continue

        if line.startswith(">"):
            flush_paragraph()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.15)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), LIGHT)
            p_pr.append(shd)
            add_inline(p, line.lstrip("> "), size=10.5, base_color=DARK_BLUE)
            index += 1
            continue

        if line.strip() == "---":
            flush_paragraph()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(6)
            paragraph_border_bottom(p, color="D8DEE7", size="4", space="2")
            index += 1
            continue

        if not line.strip():
            flush_paragraph()
        else:
            paragraph_buffer.append(line.strip())
        index += 1

    flush_paragraph()

    core = doc.core_properties
    core.title = title
    core.subject = "Báo cáo đồ án tốt nghiệp SafeFleet - Chương 1 đến Chương 3"
    core.author = "SafeFleet Project"
    core.keywords = "SafeFleet, quản lý vận hành, an toàn lái xe, OCR, AI cabin"
    core.comments = "Tài liệu tự chứa; toàn bộ sơ đồ đã được nhúng dưới dạng hình ảnh."

    output_path = Path(output_path).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Created self-contained DOCX: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) != 4:
        raise SystemExit("Usage: build_self_contained_report.py <report.md> <image-manifest.json> <output.docx>")
    build_document(sys.argv[1], sys.argv[2], sys.argv[3])
